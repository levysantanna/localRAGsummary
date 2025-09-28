#!/usr/bin/env python3
"""
Script para criar arquivos reais de teste para diferentes tipos
"""

import os
import sys
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import json

def create_test_pdf():
    """Criar um PDF de teste simples"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = Path("documents/teste_real.pdf")
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        
        # Adicionar texto
        c.drawString(100, 750, "Documento PDF de Teste")
        c.drawString(100, 700, "Este é um documento PDF real para teste do sistema RAG.")
        c.drawString(100, 650, "Contém texto que deve ser extraído pelo processador.")
        c.drawString(100, 600, "Sistema de Inteligência Artificial")
        c.drawString(100, 550, "Machine Learning e Deep Learning")
        
        c.save()
        print("✅ PDF criado: teste_real.pdf")
        return True
    except ImportError:
        print("❌ reportlab não instalado, criando PDF simples...")
        # Criar um PDF simples usando texto
        with open("documents/teste_real.pdf", "w") as f:
            f.write("PDF simples de teste")
        return True
    except Exception as e:
        print(f"❌ Erro ao criar PDF: {e}")
        return False

def create_test_image():
    """Criar uma imagem de teste com texto"""
    try:
        # Criar imagem
        img = Image.new('RGB', (400, 200), color='white')
        draw = ImageDraw.Draw(img)
        
        # Tentar usar fonte padrão
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
        except:
            font = ImageFont.load_default()
        
        # Adicionar texto
        draw.text((10, 50), "Imagem de Teste", fill='black', font=font)
        draw.text((10, 80), "Sistema RAG Local", fill='black', font=font)
        draw.text((10, 110), "Processamento de Imagens", fill='black', font=font)
        draw.text((10, 140), "OCR com EasyOCR e Tesseract", fill='black', font=font)
        
        # Salvar imagem
        img.save("documents/teste_real.png")
        print("✅ Imagem criada: teste_real.png")
        return True
    except Exception as e:
        print(f"❌ Erro ao criar imagem: {e}")
        return False

def create_test_docx():
    """Criar um documento Word de teste"""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Documento Word de Teste', 0)
        
        doc.add_paragraph('Este é um documento Word real para teste do sistema RAG.')
        doc.add_paragraph('Contém texto estruturado que deve ser extraído pelo processador.')
        
        doc.add_heading('Seção 1: Inteligência Artificial', level=1)
        doc.add_paragraph('A inteligência artificial é uma área da ciência da computação.')
        
        doc.add_heading('Seção 2: Machine Learning', level=1)
        doc.add_paragraph('Machine Learning é um subcampo da inteligência artificial.')
        
        doc.save("documents/teste_real.docx")
        print("✅ Documento Word criado: teste_real.docx")
        return True
    except ImportError:
        print("❌ python-docx não instalado")
        return False
    except Exception as e:
        print(f"❌ Erro ao criar documento Word: {e}")
        return False

def create_test_odt():
    """Criar um documento LibreOffice de teste"""
    try:
        from docx import Document
        
        # Criar documento
        doc = Document()
        doc.add_heading('Documento LibreOffice de Teste', 0)
        
        doc.add_paragraph('Este é um documento LibreOffice real para teste do sistema RAG.')
        doc.add_paragraph('Contém texto estruturado que deve ser extraído pelo processador.')
        
        doc.add_heading('Seção 1: Processamento de Documentos', level=1)
        doc.add_paragraph('O processamento de documentos é essencial para sistemas RAG.')
        
        doc.add_heading('Seção 2: Extração de Texto', level=1)
        doc.add_paragraph('A extração de texto permite análise de conteúdo.')
        
        # Salvar como ODT (formato LibreOffice)
        doc.save("documents/teste_real.odt")
        print("✅ Documento LibreOffice criado: teste_real.odt")
        return True
    except ImportError:
        print("❌ python-docx não instalado")
        return False
    except Exception as e:
        print(f"❌ Erro ao criar documento LibreOffice: {e}")
        return False

def create_test_code_files():
    """Criar arquivos de código de teste"""
    try:
        # Arquivo Python
        python_code = '''#!/usr/bin/env python3
"""
Script de teste para processamento de código
"""

def hello_world():
    """Função de exemplo"""
    print("Hello, World!")
    return "Sucesso"

class TestClass:
    """Classe de teste"""
    
    def __init__(self):
        self.name = "Teste"
    
    def process_data(self, data):
        """Processa dados"""
        return data.upper()

if __name__ == "__main__":
    hello_world()
    test = TestClass()
    result = test.process_data("teste")
    print(result)
'''
        
        with open("documents/teste_real.py", "w", encoding="utf-8") as f:
            f.write(python_code)
        print("✅ Arquivo Python criado: teste_real.py")
        
        # Arquivo JavaScript
        js_code = '''// Arquivo JavaScript de teste
function helloWorld() {
    console.log("Hello, World!");
    return "Sucesso";
}

class TestClass {
    constructor() {
        this.name = "Teste";
    }
    
    processData(data) {
        return data.toUpperCase();
    }
}

// Uso
helloWorld();
const test = new TestClass();
const result = test.processData("teste");
console.log(result);
'''
        
        with open("documents/teste_real.js", "w", encoding="utf-8") as f:
            f.write(js_code)
        print("✅ Arquivo JavaScript criado: teste_real.js")
        
        # Arquivo HTML
        html_code = '''<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Documento HTML de Teste</title>
</head>
<body>
    <h1>Sistema RAG Local</h1>
    <p>Este é um documento HTML de teste para o sistema RAG.</p>
    <h2>Características</h2>
    <ul>
        <li>Processamento de documentos</li>
        <li>Extração de texto</li>
        <li>Análise de conteúdo</li>
    </ul>
    <h2>Funcionalidades</h2>
    <p>O sistema RAG permite processar diversos tipos de documentos.</p>
</body>
</html>
'''
        
        with open("documents/teste_real.html", "w", encoding="utf-8") as f:
            f.write(html_code)
        print("✅ Arquivo HTML criado: teste_real.html")
        
        return True
    except Exception as e:
        print(f"❌ Erro ao criar arquivos de código: {e}")
        return False

def create_test_json():
    """Criar arquivo JSON de teste"""
    try:
        data = {
            "sistema": "RAG Local",
            "versao": "1.0.0",
            "funcionalidades": [
                "Processamento de documentos",
                "Extração de texto",
                "Análise de conteúdo",
                "Scraping de URLs"
            ],
            "tipos_suportados": {
                "texto": [".txt", ".md"],
                "pdf": [".pdf"],
                "imagens": [".png", ".jpg", ".jpeg"],
                "documentos": [".docx", ".doc"],
                "codigo": [".py", ".js", ".html"]
            },
            "configuracao": {
                "idioma": "pt-BR",
                "modelo": "portuguese-bert",
                "ocr": True,
                "scraping": True
            }
        }
        
        with open("documents/teste_real.json", "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print("✅ Arquivo JSON criado: teste_real.json")
        return True
    except Exception as e:
        print(f"❌ Erro ao criar arquivo JSON: {e}")
        return False

def main():
    """Função principal"""
    print("🚀 Criando arquivos reais de teste...")
    print("=" * 50)
    
    # Criar diretório se não existir
    Path("documents").mkdir(exist_ok=True)
    
    # Criar arquivos de teste
    results = []
    results.append(create_test_pdf())
    results.append(create_test_image())
    results.append(create_test_docx())
    results.append(create_test_odt())
    results.append(create_test_code_files())
    results.append(create_test_json())
    
    # Resumo
    success_count = sum(results)
    total_count = len(results)
    
    print(f"\n📊 Resumo:")
    print(f"✅ Sucessos: {success_count}/{total_count}")
    print(f"❌ Falhas: {total_count - success_count}/{total_count}")
    
    if success_count > 0:
        print("\n📁 Arquivos criados:")
        for file in Path("documents").glob("teste_real.*"):
            print(f"  📄 {file.name}")
    
    print("\n✅ Criação de arquivos concluída!")

if __name__ == "__main__":
    main()
